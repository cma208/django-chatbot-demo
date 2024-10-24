import os
import pickle
import mimetypes
import faiss
import time

from langchain_openai import OpenAIEmbeddings
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate
from langchain.chains import LLMChain
from langchain_community.vectorstores import FAISS
from langchain_community.chat_models import ChatOpenAI
from chatbot_demo.settings import *
from .models import Docs, DocThemes, ChatSession, ChatMessage
from openai import OpenAIError


from PyPDF2 import PdfReader
from docx import Document as DocxDocument

recursive_text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)

# Main logic
def get_answer_from_index_with_memory(question, index_name, session_id, max_context_messages=10):
    try:
        # Retrieve the chat session
        session = ChatSession.objects.get(session_id=session_id)
        
        # Retrieve only the most recent 'max_context_messages' from the session
        recent_messages = ChatMessage.objects.filter(session=session).order_by('-created_at')[:max_context_messages]
        
        # Reverse the order to maintain the correct conversation flow
        recent_messages = reversed(recent_messages)
        
        # Compile recent conversation history
        context_messages = list(recent_messages)
        
        # Step 1: Query the FAISS index to retrieve relevant document chunks
        retrieved_docs = query_faiss_index(question, index_name)
        
        # Step 2: Generate an answer based on the retrieved chunks and past conversation context
        answer = generate_answer(question, retrieved_docs, context_messages)
        
        # Save the new user question and answer to the database
        ChatMessage.objects.create(session=session, message=question, response=answer)
        
        return answer
    except ChatSession.DoesNotExist:
        print(f"Session with ID {session_id} not found.")
        return "Session not found."
    except Exception as e:
        print(f"An error occurred while generating the answer: {e}")
        return "I am unable to answer that question at the moment."    

# Function to load FAISS index and perform a search
def query_faiss_index(query, index_name):
    doc_path = os.path.join(FAISS_INDEX_FILE, index_name)
    
    if not os.path.exists(os.path.join(doc_path, "index.faiss")):
        raise ValueError(f"FAISS index for {index_name} does not exist.")

    # Load the existing FAISS index
    faiss_index = FAISS.load_local(doc_path, embeddings, allow_dangerous_deserialization=True)
    
    # Perform a similarity search
    retrieved_docs = faiss_index.similarity_search(query, k=5)  # Retrieve top 5 relevant chunks
    
    return retrieved_docs

def generate_answer(question, retrieved_docs, context_messages):
    # Combine the content of the retrieved docs
    doc_context = "\n\n".join([doc.page_content for doc in retrieved_docs])
    
    # Build context from past user messages
    message_context = "\n".join([
        f"User: {msg.message}\nAI: {msg.response}" for msg in context_messages
    ])
    
    # Prepare a prompt to pass to the LLM
    # The context now includes both the retrieved documents and past conversation history
    template = ChatPromptTemplate.from_template(
        template=(
            "You are a helpful chatbot assistant, not related to any doc, developed by Carlos Machiado Ardiles. Refer to the context from the past interactions below to answer the new question if needed or asked.\n"
            "Past Conversation:\n{message_context}\n\n"
            "Relevant Doc Information:\n{doc_context}\n\n"
            "Question: {question}\nAnswer:"
        )
    )
    
    chain = LLMChain(llm=llm, prompt=template, output_parser=StrOutputParser())
    
    # Generate an answer using the LLM
    answer = chain.run({"message_context": message_context, "doc_context": doc_context, "question": question})
    
    return answer

# Generate FAISS index and save it to file
def generate_faiss_with_retry(index_id, index_name):
    try:
        return generate_faiss(index_id, index_name)
    except OpenAIError as api_error:
        print(f"Error interacting with OpenAI: {api_error}")
        time.sleep(2)  # Optional: Add delay before retrying
        return generate_faiss(index_id, index_name)  # Retry once more
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return []
    
def generate_faiss(index_id,index_name):
    
    # Theme index filter
    doc_theme = DocThemes.objects.filter(id=index_id).first()

    # Check if the theme exists by ID or name, and create a new one if not found
    doc_theme = DocThemes.objects.filter(id=index_id).first()
    if not doc_theme:
        if index_name == "":
            raise ValueError("WrongIndex: The specified index does not exist.")
        else:
            print(f"Storing in new index: {index_name}")
            try:
                # Create a new theme in the database
                doc_theme = DocThemes.objects.create(theme=index_name)
            except Exception as e:
                print(f"An unexpected error occurred: {e}")
                return []

    index_id = doc_theme.id
    index_name = doc_theme.theme

    # Create a new index directory if it doesn't exist
    doc_path = os.path.join(FAISS_INDEX_FILE, index_name)
    os.makedirs(doc_path, exist_ok=True)

    docs = Docs.objects.filter(theme__id=index_id, faiss_loaded=False)
    file_names = []  # To keep track of newly added file names

    all_documents = []
    for doc in docs:
        file_path = os.path.join(BASE_DIR, doc.file.name)  # Get the actual file path
        content = robust_extract_text(file_path)  # Extract text based on file type
        
        # Add the file name to the list
        file_names.append(doc.file.name)
        
        # Create document chunks
        book_documents = recursive_text_splitter.create_documents([content])
            
        # Clean and prepare documents
        book_documents = [
            Document(page_content=text.page_content.replace("\n", " ").replace(".", "").replace("-", ""))
            for text in book_documents
        ]
        all_documents.extend(book_documents)

        if all_documents:
            # Load existing FAISS index and merge new content
            if os.path.exists(os.path.join(doc_path, "index.faiss")):
                old_docsearch = FAISS.load_local(doc_path, embeddings, allow_dangerous_deserialization=True)
                new_docsearch = FAISS.from_documents(all_documents, embeddings)
                new_docsearch.merge_from(old_docsearch)
            else:
                new_docsearch = FAISS.from_documents(all_documents, embeddings)

            # Save the updated index
            new_docsearch.save_local(doc_path)

            # Mark the processed documents as loaded in the database
            Docs.objects.filter(id__in=[doc.id for doc in docs]).update(faiss_loaded=True)

    return file_names

def robust_extract_text(file_path):
    try:
        return extract_text(file_path)
    except ValueError as e:
        print(f"Error while extracting text from {file_path}: {str(e)}")
        return "Failed to process document."
    except OpenAIError as api_error:
        print(f"OpenAI API Error: {api_error}")
        # Optional: You can set a retry or fallback action here
        return "There was an issue with the language model. Please try again later."
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        return "An unexpected error occurred. Please check the system logs."

def extract_text(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    
    if mime_type == 'application/pdf':
        # Extract text from PDF
        return extract_text_from_pdf(file_path)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':  # .docx MIME type
        # Extract text from DOCX
        return extract_text_from_docx(file_path)
    elif mime_type.startswith('text'):  # Plain text files
        # Extract text from plain text file
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        reader = PdfReader(f)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text
